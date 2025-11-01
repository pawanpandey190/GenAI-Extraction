# backend/utils.py
import pdfplumber, docx, io, re
import os
import pandas as pd
from docx import Document
from datetime import datetime
import uuid, shutil



OUTPUT_DIR = "outputs"
os.makedirs(OUTPUT_DIR, exist_ok=True)

def parse_pdf_bytes(b: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(b)) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            if t:
                text_parts.append(t)
    return "\n\n".join(text_parts).strip()

def parse_docx_bytes(b: bytes) -> str:
    doc = docx.Document(io.BytesIO(b))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)

def parse_txt_bytes(b: bytes) -> str:
    try:
        return b.decode("utf-8")
    except:
        return b.decode("latin-1", errors="ignore")

def extract_text_from_file(filename: str, content: bytes) -> str:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext == "pdf":
        return parse_pdf_bytes(content)
    elif ext == "docx":
        return parse_docx_bytes(content)
    elif ext == "txt":
        return parse_txt_bytes(content)
    else:
        raise ValueError("Unsupported file type")

def chunk_text(text: str, chunk_size=1000, overlap=200):
    sentences = re.split(r'(?<=[.!?])\s+', text)
    chunks, current = [], ""
    for s in sentences:
        if len(current)+len(s) <= chunk_size:
            current += (" " if current else "") + s
        else:
            if current: chunks.append(current.strip())
            current = s
    if current: chunks.append(current.strip())
    merged=[]
    for ch in chunks:
        if merged and len(merged[-1])+len(ch) <= chunk_size+overlap:
            merged[-1] += " " + ch
        else: merged.append(ch)
    return merged



# OUTPUT_DIR = "reports"
def merge_to_excel_and_word(all_answers,output_dir):
    """
    Merges all Q/A pairs into Excel and Word files.
    Returns full file paths for both.
    """

    # Safety checks
    if not all_answers or not isinstance(all_answers, list):
        raise ValueError("all_answers must be a non-empty list of Q/A dictionaries")

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Unique filename for each report (avoid overwriting)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    unique_id = uuid.uuid4().hex[:6]
    excel_filename = f"PPI_Report_{timestamp}_{unique_id}.csv"
    word_filename = f"PPI_Report_{timestamp}_{unique_id}.docx"

    excel_path = os.path.join(OUTPUT_DIR, excel_filename)
    word_path = os.path.join(OUTPUT_DIR, word_filename)

    # --- Excel Generation ---
    try:
        df = pd.DataFrame(all_answers)
        if df.empty:
            raise ValueError("No data to write in csv report.")
        df.to_csv(excel_path, index=False)
    except Exception as e:
        raise RuntimeError(f"Failed to generate csv report: {e}")

    # --- Word Generation ---
    try:
        doc = Document()
        doc.add_heading("PPI Parameters Extraction Report", level=1)

        for qa in all_answers:
            question = qa.get("Question") or qa.get("question") or "N/A"
            answer = qa.get("Answer") or qa.get("answer") or "N/A"
            doc.add_heading(question, level=2)
            doc.add_paragraph(answer)

        doc.save(word_path)
    except Exception as e:
        raise RuntimeError(f"Failed to generate Word report: {e}")

    return excel_path, word_path